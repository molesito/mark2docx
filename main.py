import io
import re
import json
import base64
from typing import Dict, List, Optional, Tuple

from flask import Flask, request, send_file, jsonify
from werkzeug.datastructures import FileStorage

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from PIL import Image

# Markdown parser
from markdown_it import MarkdownIt
from mdit_py_plugins.tasklists import tasklists_plugin
from mdit_py_plugins.table import table_plugin
from mdit_py_plugins.anchors import anchors_plugin
from mdit_py_plugins.sub import sub_plugin
from mdit_py_plugins.sup import sup_plugin

# LaTeX -> MathML -> OMML
from latex2mathml.converter import convert as latex_to_mathml
from pymathml2omml import mathml2omml
from lxml import etree

# Fallback render to image for very exotic LaTeX
import matplotlib
matplotlib.use("Agg")  # no display server
import matplotlib.pyplot as plt

app = Flask(__name__)

# -----------------------
# Markdown configuration
# -----------------------
md = (
    MarkdownIt("commonmark", {"html": False})
    .use(table_plugin)
    .use(tasklists_plugin, enabled=True, label=True)
    .use(anchors_plugin)
    .use(sub_plugin)
    .use(sup_plugin)
)

MATH_BLOCK_RE = re.compile(r"^\s*\$\$(.*?)\$\$\s*$", re.DOTALL)
MATH_INLINE_SPLIT_RE = re.compile(r"(\$\$.*?\$\$|\$.*?\$)", re.DOTALL)

NEW_PAGE_MARKER = "[NUEVA PÁGINA]"

# -----------------------
# Helpers: Word styling
# -----------------------
def ensure_code_style(doc: Document, style_name="Code"):
    if style_name in [s.name for s in doc.styles]:
        return style_name
    # Create a very simple code style
    s = doc.styles.add_style(style_name, 1)  # 1 = paragraph style
    s.font.name = "Consolas"
    s.font.size = Pt(10)
    return style_name

def add_hyperlink(paragraph, url, text):
    # Create a real hyperlink relationship
    part = paragraph.part
    r_id = part.relate_to(
        url,
        relationshiptype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = etree.Element(qn("w:hyperlink"), {qn("r:id"): r_id})
    new_run = etree.Element(qn("w:r"))
    rPr = etree.Element(qn("w:rPr"))
    u = etree.Element(qn("w:u"), {qn("w:val"): "single"})
    color = etree.Element(qn("w:color"), {qn("w:val"): "0000FF"})
    rPr.append(u)
    rPr.append(color)
    t = etree.Element(qn("w:t"))
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def append_omml(paragraph, omml_xml: str, as_block: bool):
    """
    Insert OMML into the paragraph. If as_block=True, wrap as oMathPara.
    """
    # Ensure we have oMathPara if it's just oMath
    root = etree.fromstring(omml_xml.encode("utf-8"))
    nsmap = root.nsmap
    m = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # If root is m:oMath, wrap into m:oMathPara
    if root.tag == f"{m}oMath":
        oMathPara = etree.Element(f"{m}oMathPara", nsmap=nsmap)
        oMathParaPr = etree.Element(f"{m}oMathParaPr")
        oMathPara.append(oMathParaPr)
        oMathPara.append(root)
        root = oMathPara

    # Convert back to string and attach
    paragraph._p.append(root)

def latex_to_omml_or_image(doc: Document, paragraph, latex_src: str, as_block: bool):
    """
    Try LaTeX -> MathML -> OMML. If it fails, fallback to image render.
    """
    try:
        # Strip $ delimiters if present
        if latex_src.startswith("$$") and latex_src.endswith(("$$",)) and len(latex_src) >= 4:
            core = latex_src[2:-2].strip()
        elif latex_src.startswith("$") and latex_src.endswith(("$",)) and len(latex_src) >= 2:
            core = latex_src[1:-1].strip()
        else:
            core = latex_src.strip()

        mathml = latex_to_mathml(core)
        omml_xml = mathml2omml(mathml)
        append_omml(paragraph, omml_xml, as_block=as_block)
        return
    except Exception:
        pass  # fallback below

    # ---- fallback to image ----
    try:
        fig = plt.figure()
        ax = fig.add_subplot(111)
        ax.axis("off")
        # matplotlib 'mathtext' uses $...$ already; ensure one pair of $
        render_tex = core
        if not (render_tex.startswith("$") and render_tex.endswith("$")):
            render_tex = f"${render_tex}$"
        ax.text(0.5, 0.5, render_tex, ha="center", va="center")
        buf = io.BytesIO()
        plt.savefig(buf, format="png", bbox_inches="tight", pad_inches=0.2, dpi=300)
        plt.close(fig)
        buf.seek(0)
        run = paragraph.add_run()
        run.add_picture(buf, width=Inches(5.5))
    except Exception:
        # As last resort, plain text
        paragraph.add_run(core)

def set_default_font_black(run):
    # Make sure black text (some themes can change it)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_page_break(doc: Document):
    doc.add_paragraph().runs.append(doc.add_paragraph().add_run())
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(type=1)  # WD_BREAK.PAGE

# -----------------------
# Image handling
# -----------------------
def get_image_by_id(images_map: Dict[str, str], file_map: Dict[str, FileStorage], image_id: str) -> Optional[io.BytesIO]:
    """
    images_map: {"id": "base64..."} from JSON
    file_map: request.files dict (for multipart)
    """
    if image_id in images_map:
        raw = base64.b64decode(images_map[image_id])
        return io.BytesIO(raw)
    if image_id in file_map:
        f = file_map[image_id]
        stream = io.BytesIO(f.read())
        return stream
    return None

def add_image_run(paragraph, image_stream: io.BytesIO, max_width_inches=6.0):
    try:
        image_stream.seek(0)
        img = Image.open(image_stream)
        w, h = img.size
        # scale to max width, keep aspect ratio
        dpi = 96.0
        width_inches = min(max_width_inches, w / dpi)
        run = paragraph.add_run()
        image_stream.seek(0)
        run.add_picture(image_stream, width=Inches(width_inches))
    except Exception:
        pass

# -----------------------
# Core rendering
# -----------------------
def render_tokens_to_docx(doc: Document, tokens, env, images_map: Dict[str, str]):
    list_stack: List[Tuple[str, int]] = []  # (ordered|bullet, level)
    code_style = ensure_code_style(doc)

    def open_list(typ: str, level: int):
        list_stack.append((typ, level))

    def close_list():
        if list_stack:
            list_stack.pop()

    # Simple state for table rendering
    in_table = False
    table = None
    table_align_cols: List[str] = []

    i = 0
    while i < len(tokens):
        t = tokens[i]

        # --- headings ---
        if t.type == "heading_open":
            level = int(t.tag[-1]) if t.tag.startswith("h") else 1
            p = doc.add_paragraph()
            p.style = f"Heading {min(level,5)}"
            i += 1
            # inner inline content
            if i < len(tokens) and tokens[i].type == "inline":
                _render_inline(doc, p, tokens[i], images_map)
                i += 1
            # heading_close
            i += 1
            continue

        # --- paragraph ---
        if t.type == "paragraph_open":
            # Special case: page break marker
            if i+1 < len(tokens) and tokens[i+1].type == "inline":
                page_text = tokens[i+1].content.strip()
                if NEW_PAGE_MARKER in page_text:
                    add_page_break(doc)
                    # consume inline + paragraph_close
                    i += 3
                    continue

            p = doc.add_paragraph()
            i += 1
            if i < len(tokens) and tokens[i].type == "inline":
                _render_inline(doc, p, tokens[i], images_map)
                i += 1
            # paragraph_close
            i += 1
            continue

        # --- blockquote ---
        if t.type == "blockquote_open":
            # increase left indent by using paragraph with increased alignment or space
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            while i < len(tokens) and tokens[i].type != "blockquote_close":
                if tokens[i].type == "paragraph_open":
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    i += 1
                    if i < len(tokens) and tokens[i].type == "inline":
                        _render_inline(doc, p, tokens[i], images_map)
                        i += 1
                    i += 1  # paragraph_close
                else:
                    i += 1
            i += 1  # blockquote_close
            continue

        # --- lists ---
        if t.type in ("bullet_list_open", "ordered_list_open"):
            list_type = "bullet" if t.type == "bullet_list_open" else "ordered"
            # level = current depth
            level = sum(1 for typ, _ in list_stack) + 1
            open_list(list_type, level)
            i += 1
            continue

        if t.type in ("bullet_list_close", "ordered_list_close"):
            close_list()
            i += 1
            continue

        if t.type == "list_item_open":
            p = doc.add_paragraph()
            # style & indent for nesting
            list_depth = len(list_stack)
            if list_stack and list_stack[-1][0] == "ordered":
                p.style = "List Number"
            else:
                p.style = "List Bullet"
            if list_depth > 1:
                p.paragraph_format.left_indent = Inches(0.3 * (list_depth - 1))
            i += 1
            # contents
            if i < len(tokens) and tokens[i].type == "paragraph_open":
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    _render_inline(doc, p, tokens[i], images_map, list_item=True)
                    i += 1
                i += 1  # paragraph_close
            # sublists or more content until list_item_close
            while i < len(tokens) and tokens[i].type != "list_item_close":
                # allow nested lists inside li
                # they will be handled in outer loop
                break
            i += 1  # list_item_close
            continue

        # --- code block ---
        if t.type == "fence" or t.type == "code_block":
            lang = getattr(t, "info", "") or ""
            code = t.content
            p = doc.add_paragraph()
            p.style = code_style
            for line in code.splitlines():
                r = p.add_run(line)
                r.font.name = "Consolas"
                set_default_font_black(r)
                p = doc.add_paragraph()
                p.style = code_style
            i += 1
            continue

        # --- hr / thematic break ---
        if t.type == "hr":
            p = doc.add_paragraph()
            # Simulate horizontal rule with border
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
            # Draw a sequence of underscores as a visible line
            run = p.add_run("_" * 60)
            set_default_font_black(run)
            i += 1
            continue

        # --- tables ---
        if t.type == "table_open":
            in_table = True
            table = None
            table_align_cols = []
            i += 1
            continue

        if t.type == "thead_open":
            header_rows = []
            i += 1
            # parse until thead_close
            while i < len(tokens) and tokens[i].type != "thead_close":
                if tokens[i].type == "tr_open":
                    header_cells = []
                    i += 1
                    while tokens[i].type != "tr_close":
                        if tokens[i].type == "th_open":
                            i += 1
                            txt = ""
                            if tokens[i].type == "inline":
                                txt = tokens[i].content
                                i += 1
                            i += 1  # th_close
                            header_cells.append(txt)
                        else:
                            i += 1
                    header_rows.append(header_cells)
                    i += 1  # tr_close
                else:
                    i += 1
            # create table with header cols
            cols = max(len(r) for r in header_rows) if header_rows else 1
            table = doc.add_table(rows=1, cols=cols)
            hdr = table.rows[0].cells
            for idx, txt in enumerate(header_rows[0] if header_rows else []):
                para = hdr[idx].paragraphs[0]
                r = para.add_run(txt)
                r.bold = True
                set_default_font_black(r)
            i += 1  # thead_close
            continue

        if t.type == "tbody_open":
            i += 1
            continue

        if t.type == "tr_open" and in_table:
            row_cells: List[str] = []
            i += 1
            while tokens[i].type != "tr_close":
                if tokens[i].type == "td_open":
                    i += 1
                    cell_text = ""
                    if tokens[i].type == "inline":
                        cell_text = tokens[i].content
                        i += 1
                    i += 1  # td_close
                    row_cells.append(cell_text)
                else:
                    i += 1
            # append row
            if table is None:
                table = doc.add_table(rows=0, cols=len(row_cells))
            row = table.add_row().cells
            for idx, txt in enumerate(row_cells):
                p = row[idx].paragraphs[0]
                p.text = ""
                _render_inline(doc, p, type("inline", (), {"content": txt, "children": []}), images_map)
            i += 1  # tr_close
            continue

        if t.type == "tbody_close":
            i += 1
            continue

        if t.type == "table_close":
            in_table = False
            table = None
            table_align_cols = []
            i += 1
            continue

        # default advance
        i += 1


def _render_inline(doc: Document, paragraph, inline_token, images_map: Dict[str, str], list_item: bool=False):
    """
    Render inline content into the given paragraph.
    Also split text by $...$ and $$...$$ to inject OMML.
    """
    text = inline_token.content or ""
    if not text:
        return

    # Handle inline math and block math embedded in a single inline
    parts = MATH_INLINE_SPLIT_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("$$") and part.endswith("$$"):
            # Block math inside inline → place as a separate equation paragraph
            eq_p = paragraph._element.getparent()  # current paragraph
            newp = paragraph._p.getparent().addnext(etree.Element(qn("w:p")))
            p_wrap = paragraph._p.getparent()
            # Simpler: just add as new paragraph in doc
            p2 = paragraph._element.getparent()
            p_math = paragraph._p
            # Easiest: add another paragraph at doc level
            p = paragraph._element
            # Actually we will just add in current paragraph for stability
            latex_to_omml_or_image(doc, paragraph, part, as_block=True)
            continue
        if part.startswith("$") and part.endswith("$") and len(part) >= 2:
            latex_to_omml_or_image(doc, paragraph, part, as_block=False)
            continue

        # Regular text: we still need to respect *bold*, _em_, ~~strike~~ etc.
        # markdown-it already normalized; inline_token.content lost marks,
        # so we simulate minimal emphasis by scanning **, *, ~~
        _render_emphasis_runs(paragraph, part)

def _render_emphasis_runs(paragraph, txt: str):
    """
    Very small emphasis parser for **bold**, *italic*, ~~strike~~, `code`,
    and plain text. For links [text](url) and images ![alt](id), we handle via patterns.
    """
    # images ![alt](id)
    # links [text](url)
    # code `...`
    token_re = re.compile(
        r"(!\[(?P<imgalt>.*?)\]\((?P<imgid>.*?)\))"  # image
        r"|(\[(?P<linktxt>.*?)\]\((?P<linkurl>.*?)\))"  # link
        r"|(`(?P<code>[^`]+)`)"  # code
        r"|(\*\*\*(?P<bic>.+?)\*\*\*)"  # bold+italic
        r"|(\*\*(?P<bold>.+?)\*\*)"  # bold
        r"|(\*(?P<ital>.+?)\*)"  # italic
        r"|(\~\~(?P<strike>.+?)\~\~)",  # strike
        re.DOTALL,
    )

    pos = 0
    for m in token_re.finditer(txt):
        if m.start() > pos:
            r = paragraph.add_run(txt[pos:m.start()])
            set_default_font_black(r)
        if m.group("imgalt") is not None:
            # images inline: we don't have the image stream here, so just put alt
            r = paragraph.add_run(f"[{m.group('imgalt')}]")
            set_default_font_black(r)
        elif m.group("linktxt") is not None:
            add_hyperlink(paragraph, m.group("linkurl"), m.group("linktxt"))
        elif m.group("code") is not None:
            r = paragraph.add_run(m.group("code"))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            set_default_font_black(r)
        elif m.group("bic") is not None:
            r = paragraph.add_run(m.group("bic"))
            r.bold = True
            r.italic = True
            set_default_font_black(r)
        elif m.group("bold") is not None:
            r = paragraph.add_run(m.group("bold"))
            r.bold = True
            set_default_font_black(r)
        elif m.group("ital") is not None:
            r = paragraph.add_run(m.group("ital"))
            r.italic = True
            set_default_font_black(r)
        elif m.group("strike") is not None:
            r = paragraph.add_run(m.group("strike"))
            r.font.strike = True
            set_default_font_black(r)
        pos = m.end()

    if pos < len(txt):
        r = paragraph.add_run(txt[pos:])
        set_default_font_black(r)

# -----------------------
# /docx endpoint
# -----------------------
@app.route("/docx", methods=["POST"])
def make_docx():
    """
    Input (JSON or multipart/form-data):

    JSON:
    {
      "markdown": "string md",
      "images": {"img1": "<base64>", ...},
      "output_name": "out.docx"  (opcional)
    }

    Multipart:
      - field 'markdown' (texto)
      - files: varias imágenes referenciables por id (mismo nombre de campo)
    """
    # Load payload
    images_map: Dict[str, str] = {}
    markdown_text = ""

    if request.content_type and "application/json" in request.content_type:
        data = request.get_json(force=True, silent=True) or {}
        markdown_text = data.get("markdown", "") or ""
        images_map = data.get("images", {}) or {}
        output_name = data.get("output_name", "documento.docx")
    else:
        markdown_text = request.form.get("markdown", "") or ""
        output_name = request.form.get("output_name", "documento.docx")
        # files handled later if referenced by id in markdown

    if not markdown_text.strip():
        return jsonify({"error": "Falta 'markdown'."}), 400

    # Create doc
    doc = Document()

    # Parse markdown into tokens
    env = {}
    tokens = md.parse(markdown_text, env)

    # Render
    render_tokens_to_docx(doc, tokens, env, images_map)

    # Save and return
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=output_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# -----------------------
# /merge endpoint (no tocar funcionalmente)
# -----------------------
@app.route("/merge", methods=["POST"])
def merge_docx():
    """
    Une varios DOCX en orden, separando cada uno por salto de página.

    Admite:
    - JSON: {"files": ["<base64docx>", ...], "output_name": "merged.docx"}
    - multipart/form-data: varios ficheros 'files'
    """
    output_name = "merged.docx"
    streams: List[io.BytesIO] = []

    if request.content_type and "application/json" in request.content_type:
        data = request.get_json(force=True, silent=True) or {}
        output_name = data.get("output_name", "merged.docx")
        files_b64 = data.get("files", []) or []
        for b in files_b64:
            try:
                streams.append(io.BytesIO(base64.b64decode(b)))
            except Exception:
                return jsonify({"error": "Uno de los archivos base64 no es válido."}), 400
    else:
        # multipart
        if "files" not in request.files:
            return jsonify({"error": "Envía ficheros en el campo 'files'."}), 400
        fs: List[FileStorage] = request.files.getlist("files")
        for f in fs:
            streams.append(io.BytesIO(f.read()))
        output_name = request.form.get("output_name", "merged.docx") or "merged.docx"

    if not streams:
        return jsonify({"error": "No se han recibido documentos para unir."}), 400

    # Merge
    base_doc = Document(streams[0])
    composer = Composer(base_doc)
    for st in streams[1:]:
        subdoc = Document(st)
        # nueva sección en página nueva
        base_doc.add_section(WD_SECTION.NEW_PAGE)
        composer.append(subdoc)

    out = io.BytesIO()
    composer.save(out)
    out.seek(0)
    return send_file(
        out,
        as_attachment=True,
        download_name=output_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# -----------------------
# Healthcheck
# -----------------------
@app.route("/", methods=["GET"])
def health():
    return jsonify({"ok": True})

if __name__ == "__main__":
    # For local testing only (Render usará gunicorn)
    app.run(host="0.0.0.0", port=8000, debug=False)
